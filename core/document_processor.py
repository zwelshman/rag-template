"""
Document Processor
Handles loading, parsing, and chunking of various document types.
Consolidates document_loader and text_splitter functionality.
"""

import os
import json
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
import pandas as pd
import re

# Supported file extensions
SUPPORTED_EXTENSIONS = ['.txt', '.pdf', '.docx', '.doc', '.xlsx', '.xls', '.csv', '.json']


@dataclass
class Document:
    """Represents a loaded document with content and metadata."""
    content: str
    metadata: Dict[str, Any]

    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}


@dataclass
class TextChunk:
    """Represents a chunk of text with metadata."""
    content: str
    metadata: dict
    chunk_index: int

    @property
    def id(self) -> str:
        """Generate a unique ID for the chunk."""
        source = self.metadata.get('source', 'unknown')
        return f"{source}_chunk_{self.chunk_index}"


class DocumentProcessor:
    """
    Universal document processor that handles:
    - Loading documents from various formats
    - Splitting documents into chunks
    """

    SUPPORTED_EXTENSIONS_MAP = {
        '.txt': 'text',
        '.pdf': 'pdf',
        '.docx': 'word',
        '.doc': 'word',
        '.xlsx': 'excel',
        '.xls': 'excel',
        '.csv': 'csv',
        '.json': 'json',
    }

    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
    ):
        """
        Initialize document processor.

        Args:
            chunk_size: Maximum chunk size in characters
            chunk_overlap: Overlap between chunks
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def load_from_bytes(self, file_bytes: bytes, filename: str) -> List[Document]:
        """
        Load a document from bytes.

        Args:
            file_bytes: Raw bytes of the file
            filename: Original filename

        Returns:
            List of Document objects
        """
        import tempfile

        ext = os.path.splitext(filename)[1].lower()

        if ext not in self.SUPPORTED_EXTENSIONS_MAP:
            raise ValueError(f"Unsupported file type: {ext}")

        # Write bytes to temporary file and load
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp_file:
            tmp_file.write(file_bytes)
            tmp_path = tmp_file.name

        try:
            documents = self._load_file(tmp_path, ext)
            # Update metadata with original filename
            for doc in documents:
                doc.metadata['source'] = filename
            return documents
        finally:
            os.unlink(tmp_path)

    def _load_file(self, file_path: str, ext: str) -> List[Document]:
        """Load a file based on its extension."""
        file_type = self.SUPPORTED_EXTENSIONS_MAP[ext]

        if file_type == 'text':
            return self._load_text(file_path)
        elif file_type == 'pdf':
            return self._load_pdf(file_path)
        elif file_type == 'word':
            return self._load_word(file_path)
        elif file_type == 'excel':
            return self._load_excel(file_path)
        elif file_type == 'csv':
            return self._load_csv(file_path)
        elif file_type == 'json':
            return self._load_json(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    def _load_text(self, file_path: str) -> List[Document]:
        """Load a plain text file."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()

        return [Document(
            content=content,
            metadata={
                'source': os.path.basename(file_path),
                'file_type': 'text',
                'file_path': file_path,
            }
        )]

    def _load_pdf(self, file_path: str) -> List[Document]:
        """Load a PDF file."""
        from pypdf import PdfReader

        reader = PdfReader(file_path)
        documents = []

        for page_num, page in enumerate(reader.pages):
            text = page.extract_text()
            if text.strip():
                documents.append(Document(
                    content=text,
                    metadata={
                        'source': os.path.basename(file_path),
                        'file_type': 'pdf',
                        'file_path': file_path,
                        'page_number': page_num + 1,
                        'total_pages': len(reader.pages),
                    }
                ))

        if not documents:
            documents.append(Document(
                content="",
                metadata={
                    'source': os.path.basename(file_path),
                    'file_type': 'pdf',
                    'file_path': file_path,
                    'note': 'No text could be extracted from this PDF',
                }
            ))

        return documents

    def _load_word(self, file_path: str) -> List[Document]:
        """Load a Word document."""
        from docx import Document as DocxDocument

        doc = DocxDocument(file_path)

        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        content = '\n\n'.join(paragraphs)

        return [Document(
            content=content,
            metadata={
                'source': os.path.basename(file_path),
                'file_type': 'word',
                'file_path': file_path,
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables),
            }
        )]

    def _load_excel(self, file_path: str) -> List[Document]:
        """Load an Excel file."""
        excel_file = pd.ExcelFile(file_path)
        documents = []

        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            content = self._dataframe_to_text(df, sheet_name)

            if content.strip():
                documents.append(Document(
                    content=content,
                    metadata={
                        'source': os.path.basename(file_path),
                        'file_type': 'excel',
                        'file_path': file_path,
                        'sheet_name': sheet_name,
                        'rows': len(df),
                        'columns': len(df.columns),
                    }
                ))

        return documents

    def _load_csv(self, file_path: str) -> List[Document]:
        """Load a CSV file."""
        encodings = ['utf-8', 'latin-1', 'cp1252']
        df = None

        for encoding in encodings:
            try:
                df = pd.read_csv(file_path, encoding=encoding)
                break
            except UnicodeDecodeError:
                continue

        if df is None:
            raise ValueError(f"Could not decode CSV file with any of these encodings: {encodings}")

        content = self._dataframe_to_text(df)

        return [Document(
            content=content,
            metadata={
                'source': os.path.basename(file_path),
                'file_type': 'csv',
                'file_path': file_path,
                'rows': len(df),
                'columns': len(df.columns),
            }
        )]

    def _load_json(self, file_path: str) -> List[Document]:
        """Load a JSON file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        content = self._json_to_text(data)

        return [Document(
            content=content,
            metadata={
                'source': os.path.basename(file_path),
                'file_type': 'json',
                'file_path': file_path,
            }
        )]

    def _dataframe_to_text(self, df: pd.DataFrame, sheet_name: Optional[str] = None) -> str:
        """Convert a DataFrame to readable text format."""
        lines = []

        if sheet_name:
            lines.append(f"Sheet: {sheet_name}")
            lines.append("-" * 40)

        lines.append("Columns: " + ", ".join(str(col) for col in df.columns))
        lines.append("")

        for idx, row in df.iterrows():
            row_parts = []
            for col in df.columns:
                value = row[col]
                if pd.notna(value):
                    row_parts.append(f"{col}: {value}")
            if row_parts:
                lines.append(f"Row {idx + 1}: " + " | ".join(row_parts))

        return '\n'.join(lines)

    def _json_to_text(self, data: Any, indent: int = 0) -> str:
        """Convert JSON data to readable text format."""
        lines = []
        indent_str = "  " * indent

        if isinstance(data, dict):
            for key, value in data.items():
                if isinstance(value, (dict, list)):
                    lines.append(f"{indent_str}{key}:")
                    lines.append(self._json_to_text(value, indent + 1))
                else:
                    lines.append(f"{indent_str}{key}: {value}")
        elif isinstance(data, list):
            for idx, item in enumerate(data):
                if isinstance(item, (dict, list)):
                    lines.append(f"{indent_str}[{idx}]:")
                    lines.append(self._json_to_text(item, indent + 1))
                else:
                    lines.append(f"{indent_str}[{idx}]: {item}")
        else:
            lines.append(f"{indent_str}{data}")

        return '\n'.join(lines)

    def split_text(self, text: str) -> List[str]:
        """
        Split text into chunks.

        Args:
            text: The text to split

        Returns:
            List of text chunks
        """
        if not text or not text.strip():
            return []

        if len(text) <= self.chunk_size:
            return [text.strip()]

        separators = ["\n\n", "\n", ". ", "! ", "? ", "; ", ", ", " ", ""]
        return self._recursive_split(text, separators)

    def _recursive_split(self, text: str, separators: List[str]) -> List[str]:
        """Recursively split text using available separators."""
        final_chunks = []

        separator = separators[-1]
        for sep in separators:
            if sep == "" or sep in text:
                separator = sep
                break

        if separator:
            splits = text.split(separator)
        else:
            splits = list(text)

        current_chunk = []
        current_length = 0

        for split in splits:
            split_length = len(split)

            if split_length > self.chunk_size:
                if current_chunk:
                    chunk_text = separator.join(current_chunk)
                    if chunk_text.strip():
                        final_chunks.append(chunk_text.strip())
                    current_chunk = []
                    current_length = 0

                if separator in separators:
                    idx = separators.index(separator)
                    if idx < len(separators) - 1:
                        sub_chunks = self._recursive_split(split, separators[idx + 1:])
                        final_chunks.extend(sub_chunks)
                    else:
                        for i in range(0, len(split), self.chunk_size - self.chunk_overlap):
                            final_chunks.append(split[i:i + self.chunk_size].strip())
                continue

            potential_length = current_length + split_length
            if current_chunk:
                potential_length += len(separator)

            if potential_length > self.chunk_size and current_chunk:
                chunk_text = separator.join(current_chunk)
                if chunk_text.strip():
                    final_chunks.append(chunk_text.strip())

                overlap_chunks = []
                overlap_length = 0
                for chunk in reversed(current_chunk):
                    chunk_len = len(chunk)
                    if overlap_length + chunk_len <= self.chunk_overlap:
                        overlap_chunks.insert(0, chunk)
                        overlap_length += chunk_len + len(separator)
                    else:
                        break

                current_chunk = overlap_chunks
                current_length = overlap_length

            current_chunk.append(split)
            current_length += split_length + (len(separator) if current_chunk else 0)

        if current_chunk:
            chunk_text = separator.join(current_chunk)
            if chunk_text.strip():
                final_chunks.append(chunk_text.strip())

        return final_chunks

    def split_documents(self, documents: List[Document]) -> List[TextChunk]:
        """
        Split a list of documents into chunks.

        Args:
            documents: List of Document objects

        Returns:
            List of TextChunk objects
        """
        all_chunks = []
        chunk_index = 0

        for doc in documents:
            text_chunks = self.split_text(doc.content)

            for i, chunk_text in enumerate(text_chunks):
                chunk = TextChunk(
                    content=chunk_text,
                    metadata={
                        **doc.metadata,
                        'chunk_index_in_doc': i,
                        'total_chunks_in_doc': len(text_chunks),
                    },
                    chunk_index=chunk_index,
                )
                all_chunks.append(chunk)
                chunk_index += 1

        return all_chunks
