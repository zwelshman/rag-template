"""
Document Loader Module
Handles loading and extracting text from various file formats:
- Text files (.txt)
- PDF files (.pdf)
- Word documents (.docx)
- Excel files (.xlsx, .xls)
- CSV files (.csv)
"""

import os
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
import pandas as pd


@dataclass
class Document:
    """Represents a loaded document with content and metadata."""
    content: str
    metadata: Dict[str, Any]

    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}


class DocumentLoader:
    """Universal document loader supporting multiple file formats."""

    SUPPORTED_EXTENSIONS = {
        '.txt': 'text',
        '.pdf': 'pdf',
        '.docx': 'word',
        '.doc': 'word',
        '.xlsx': 'excel',
        '.xls': 'excel',
        '.csv': 'csv',
    }

    def __init__(self):
        self._loaders = {
            'text': self._load_text,
            'pdf': self._load_pdf,
            'word': self._load_word,
            'excel': self._load_excel,
            'csv': self._load_csv,
        }

    def load(self, file_path: str) -> List[Document]:
        """
        Load a document from the given file path.

        Args:
            file_path: Path to the file to load

        Returns:
            List of Document objects
        """
        ext = os.path.splitext(file_path)[1].lower()

        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {ext}. Supported types: {list(self.SUPPORTED_EXTENSIONS.keys())}")

        file_type = self.SUPPORTED_EXTENSIONS[ext]
        loader = self._loaders[file_type]

        return loader(file_path)

    def load_from_bytes(self, file_bytes: bytes, filename: str) -> List[Document]:
        """
        Load a document from bytes (useful for Streamlit file uploads).

        Args:
            file_bytes: Raw bytes of the file
            filename: Original filename (used to determine file type)

        Returns:
            List of Document objects
        """
        import tempfile

        ext = os.path.splitext(filename)[1].lower()

        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {ext}")

        # Write bytes to temporary file and load
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp_file:
            tmp_file.write(file_bytes)
            tmp_path = tmp_file.name

        try:
            documents = self.load(tmp_path)
            # Update metadata with original filename
            for doc in documents:
                doc.metadata['source'] = filename
            return documents
        finally:
            os.unlink(tmp_path)

    def _load_text(self, file_path: str) -> List[Document]:
        """Load a plain text file."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()

        return [Document(
            content=content,
            metadata={
                'source': os.path.basename(file_path),
                'file_type': 'text',
                'file_path': file_path,
            }
        )]

    def _load_pdf(self, file_path: str) -> List[Document]:
        """Load a PDF file."""
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ImportError("pypdf is required for PDF support. Install with: pip install pypdf")

        reader = PdfReader(file_path)
        documents = []

        for page_num, page in enumerate(reader.pages):
            text = page.extract_text()
            if text.strip():
                documents.append(Document(
                    content=text,
                    metadata={
                        'source': os.path.basename(file_path),
                        'file_type': 'pdf',
                        'file_path': file_path,
                        'page_number': page_num + 1,
                        'total_pages': len(reader.pages),
                    }
                ))

        # If no text was extracted, return empty document with metadata
        if not documents:
            documents.append(Document(
                content="",
                metadata={
                    'source': os.path.basename(file_path),
                    'file_type': 'pdf',
                    'file_path': file_path,
                    'note': 'No text could be extracted from this PDF',
                }
            ))

        return documents

    def _load_word(self, file_path: str) -> List[Document]:
        """Load a Word document (.docx)."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("python-docx is required for Word document support. Install with: pip install python-docx")

        doc = DocxDocument(file_path)

        # Extract text from paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        content = '\n\n'.join(paragraphs)

        return [Document(
            content=content,
            metadata={
                'source': os.path.basename(file_path),
                'file_type': 'word',
                'file_path': file_path,
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables),
            }
        )]

    def _load_excel(self, file_path: str) -> List[Document]:
        """Load an Excel file (.xlsx, .xls)."""
        try:
            import openpyxl  # noqa: F401 - verify import
        except ImportError:
            raise ImportError("openpyxl is required for Excel support. Install with: pip install openpyxl")

        # Read all sheets
        excel_file = pd.ExcelFile(file_path)
        documents = []

        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)

            # Convert DataFrame to readable text
            content = self._dataframe_to_text(df, sheet_name)

            if content.strip():
                documents.append(Document(
                    content=content,
                    metadata={
                        'source': os.path.basename(file_path),
                        'file_type': 'excel',
                        'file_path': file_path,
                        'sheet_name': sheet_name,
                        'rows': len(df),
                        'columns': len(df.columns),
                    }
                ))

        return documents

    def _load_csv(self, file_path: str) -> List[Document]:
        """Load a CSV file."""
        # Try different encodings
        encodings = ['utf-8', 'latin-1', 'cp1252']
        df = None

        for encoding in encodings:
            try:
                df = pd.read_csv(file_path, encoding=encoding)
                break
            except UnicodeDecodeError:
                continue

        if df is None:
            raise ValueError(f"Could not decode CSV file with any of these encodings: {encodings}")

        content = self._dataframe_to_text(df)

        return [Document(
            content=content,
            metadata={
                'source': os.path.basename(file_path),
                'file_type': 'csv',
                'file_path': file_path,
                'rows': len(df),
                'columns': len(df.columns),
            }
        )]

    def _dataframe_to_text(self, df: pd.DataFrame, sheet_name: Optional[str] = None) -> str:
        """Convert a DataFrame to readable text format."""
        lines = []

        if sheet_name:
            lines.append(f"Sheet: {sheet_name}")
            lines.append("-" * 40)

        # Add column headers
        lines.append("Columns: " + ", ".join(str(col) for col in df.columns))
        lines.append("")

        # Convert each row to text
        for idx, row in df.iterrows():
            row_parts = []
            for col in df.columns:
                value = row[col]
                if pd.notna(value):
                    row_parts.append(f"{col}: {value}")
            if row_parts:
                lines.append(f"Row {idx + 1}: " + " | ".join(row_parts))

        return '\n'.join(lines)

    @classmethod
    def get_supported_extensions(cls) -> List[str]:
        """Get list of supported file extensions."""
        return list(cls.SUPPORTED_EXTENSIONS.keys())
